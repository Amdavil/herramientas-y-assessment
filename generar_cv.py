# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import copy, io as _io, os

# ── RUTAS ──────────────────────────────────────────────────────────────────
PHOTO_SRC  = r"C:\Users\danie\Downloads\ChatGPT Image 7 may 2026, 10_20_26.png"
OUTPUT     = r"G:\Mi unidad\5-DOCUMENTOS DANIEL\Resume\Daniel_Villa_CV_Profesional.docx"

# ── COLORES ────────────────────────────────────────────────────────────────
NAVY       = RGBColor(0x1C, 0x35, 0x53)
NAVY_HEX   = "1C3553"
GREEN      = RGBColor(0x27, 0x96, 0x5C)
GREEN_HEX  = "27965C"
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
WHITE_HEX  = "FFFFFF"
LBLUE      = RGBColor(0xA8, 0xC8, 0xE8)
LBLUE_HEX  = "A8C8E8"
DARK       = RGBColor(0x2C, 0x2C, 0x2C)
GRAY       = RGBColor(0x77, 0x77, 0x77)
SIDEBAR_HEX= NAVY_HEX

# ── HELPERS XML ────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'),    str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)

def set_cell_valign(cell, align='top'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)

def remove_table_borders(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'none')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)

def remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'none')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tcBorders.append(el)
    tcPr.append(tcBorders)

def set_para_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'),  str(after))
    if line:
        spacing.set(qn('w:line'),     str(line))
        spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def add_bottom_border(para, hex_color, size=12):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    str(size))
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), hex_color)
    pBdr.append(bot)
    pPr.append(pBdr)

def run_color(run, hex_color):
    rPr = run._r.get_or_add_rPr()
    color = OxmlElement('w:color')
    color.set(qn('w:val'), hex_color)
    rPr.append(color)

# ── FUNCIONES DE CONTENIDO ─────────────────────────────────────────────────
def sidebar_section(cell, text):
    p = cell.add_paragraph()
    set_para_spacing(p, before=160, after=60)
    add_bottom_border(p, GREEN_HEX, size=8)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.color.rgb = LBLUE
    return p

def sidebar_line(cell, label, value, label_bold=True):
    p = cell.add_paragraph()
    set_para_spacing(p, before=30, after=30)
    if label:
        r1 = p.add_run(label)
        r1.bold = label_bold
        r1.font.name = 'Calibri'
        r1.font.size = Pt(8.5)
        r1.font.color.rgb = WHITE
    if value:
        r2 = p.add_run(value)
        r2.font.name = 'Calibri'
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)
    return p

def sidebar_item(cell, text, size=8.5):
    p = cell.add_paragraph()
    set_para_spacing(p, before=28, after=28)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)
    return p

def main_section(cell, text):
    p = cell.add_paragraph()
    set_para_spacing(p, before=220, after=80)
    add_bottom_border(p, GREEN_HEX, size=14)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    return p

def job_header(cell, company, location, dates, role):
    # Company + Location
    p1 = cell.add_paragraph()
    set_para_spacing(p1, before=180, after=20)
    r1 = p1.add_run(company)
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = NAVY
    r2 = p1.add_run(f'  |  {location}')
    r2.font.name = 'Calibri'
    r2.font.size = Pt(9)
    r2.font.color.rgb = GRAY
    r2.italic = True
    # Dates
    p2 = cell.add_paragraph()
    set_para_spacing(p2, before=0, after=30)
    r3 = p2.add_run(dates)
    r3.font.name = 'Calibri'
    r3.font.size = Pt(8.5)
    r3.font.color.rgb = GREEN
    r3.bold = True
    # Role
    p3 = cell.add_paragraph()
    set_para_spacing(p3, before=0, after=40)
    r4 = p3.add_run(role)
    r4.font.name = 'Calibri'
    r4.font.size = Pt(9.5)
    r4.font.color.rgb = DARK
    r4.italic = True
    r4.bold = True

def bullet_item(cell, text, indent_left=300, size=9):
    p = cell.add_paragraph()
    set_para_spacing(p, before=28, after=28)
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'),    str(indent_left))
    ind.set(qn('w:hanging'), '180')
    pPr.append(ind)
    dot = p.add_run('› ')
    dot.font.name  = 'Calibri'
    dot.font.size  = Pt(size)
    dot.font.color.rgb = GREEN
    dot.bold = True
    body = p.add_run(text)
    body.font.name  = 'Calibri'
    body.font.size  = Pt(size)
    body.font.color.rgb = DARK

def sub_label(cell, text):
    p = cell.add_paragraph()
    set_para_spacing(p, before=60, after=20)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.color.rgb = DARK

def main_para(cell, parts, before=40, after=40, justify=True):
    p = cell.add_paragraph()
    set_para_spacing(p, before=before, after=after)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, bold in parts:
        r = p.add_run(text)
        r.bold = bold
        r.font.name = 'Calibri'
        r.font.size = Pt(9.5)
        r.font.color.rgb = DARK
    return p

# ── IMAGEN CUADRADA CENTRADA ───────────────────────────────────────────────
def add_centered_image(cell, path, size_cm=4.0):
    img = Image.open(path)
    w, h = img.size
    side = min(w, h)
    left = (w - side) // 2
    top  = (h - side) // 2
    img_cropped = img.crop((left, top, left + side, top + side))
    buf = _io.BytesIO()
    img_cropped.save(buf, format='PNG')
    buf.seek(0)

    p = cell.add_paragraph()
    set_para_spacing(p, before=0, after=120)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Cm(size_cm))
    return p

# ══════════════════════════════════════════════════════════════════════════
#  DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════
doc = Document()

# Márgenes de página
section = doc.sections[0]
section.page_width  = Emu(12240 * 914)   # 8.5"  en EMU
section.page_height = Emu(15840 * 914)   # 11"
section.top_margin    = Cm(1.0)
section.bottom_margin = Cm(1.0)
section.left_margin   = Cm(0.7)
section.right_margin  = Cm(0.7)

# Tabla raíz: 1 fila, 2 columnas
SIDEBAR_W_CM = 6.5   # cm
MAIN_W_CM    = 13.2  # cm

table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
remove_table_borders(table)

# Ancho de columnas (en DXA: 1 cm ≈ 567 DXA)
tbl = table._tbl
tblGrid = OxmlElement('w:tblGrid')
for w_cm in [SIDEBAR_W_CM, MAIN_W_CM]:
    gridCol = OxmlElement('w:gridCol')
    gridCol.set(qn('w:w'), str(int(w_cm * 567)))
    tblGrid.append(gridCol)
tbl.insert(1, tblGrid)

# Ancho total de la tabla
tblPr = tbl.find(qn('w:tblPr'))
if tblPr is None:
    tblPr = OxmlElement('w:tblPr')
    tbl.insert(0, tblPr)
tblW = OxmlElement('w:tblW')
tblW.set(qn('w:w'),    str(int((SIDEBAR_W_CM + MAIN_W_CM) * 567)))
tblW.set(qn('w:type'), 'dxa')
tblPr.append(tblW)

row  = table.rows[0]
left = row.cells[0]
right= row.cells[1]

# Ancho individual de celdas
for cell, w_cm in [(left, SIDEBAR_W_CM), (right, MAIN_W_CM)]:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'),    str(int(w_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)
    remove_cell_borders(cell)

set_cell_bg(left,  SIDEBAR_HEX)
set_cell_bg(right, WHITE_HEX)
set_cell_margins(left,  top=200, bottom=200, left=220, right=220)
set_cell_margins(right, top=200, bottom=200, left=280, right=180)
set_cell_valign(left,  'top')
set_cell_valign(right, 'top')

# ══════════════════════════════════════════════════════════════════════════
#  SIDEBAR  (columna izquierda)
# ══════════════════════════════════════════════════════════════════════════

# — Foto
add_centered_image(left, PHOTO_SRC, size_cm=3.8)

# — Nombre
p_name = left.add_paragraph()
set_para_spacing(p_name, before=0, after=30)
p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
rn = p_name.add_run('DANIEL VILLA VÉLEZ')
rn.bold = True
rn.font.name = 'Calibri'
rn.font.size = Pt(13)
rn.font.color.rgb = WHITE

# — Título
p_title = left.add_paragraph()
set_para_spacing(p_title, before=0, after=180)
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = p_title.add_run('Ingeniero Ambiental  |  MBA')
rt.font.name = 'Calibri'
rt.font.size = Pt(8.5)
rt.font.color.rgb = LBLUE
rt.italic = True

# — CONTACTO
sidebar_section(left, 'Contacto')
sidebar_line(left, 'Tel:  ',    '+57 300 235 4198')
sidebar_line(left, 'Email:  ',  'danielvillavelez@gmail.com')
sidebar_line(left, 'Ciudad:  ', 'Envigado, Antioquia')
sidebar_line(left, 'Web:  ',    'bit.ly/2RRr0pe')

# — IDIOMAS
sidebar_section(left, 'Idiomas')
sidebar_line(left, 'Español  ',  'Lengua materna')
sidebar_line(left, 'Inglés  ',   'Alto')
sidebar_line(left, 'Francés  ',  'Básico')

# — ÁREAS DE EXPERTICIA
sidebar_section(left, 'Áreas de Experticia')
for skill in [
    'Planeación estratégica y proyectos',
    'Sostenibilidad corporativa y ESG',
    'Reportes GRI / estándares ISSB',
    'Cambio climático y huella de carbono',
    'Transformación digital',
    'IA generativa (ChatGPT, Claude, Copilot)',
    'Automatización de procesos con IA',
    'Responsabilidad social corporativa',
]:
    sidebar_item(left, f'› {skill}')

# — EDUCACIÓN FORMAL
sidebar_section(left, 'Educación Formal')
for degree, school, year in [
    ('MBA Ejecutivo',                   'Westfield Business School',             '2021 – 2022'),
    ('Máster en Transformación Exp.',   'EIG – Granada, España',                 '2022'),
    ('Esp. Gerencia de Proyectos',      'Univ. Pontificia Bolivariana',          '2006 – 2008'),
    ('Ingeniero Ambiental',             'Escuela de Ing. de Antioquia',          '1998 – 2004'),
]:
    p = left.add_paragraph()
    set_para_spacing(p, before=100, after=0)
    r1 = p.add_run(degree)
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(8.5)
    r1.font.color.rgb = WHITE
    p2 = left.add_paragraph()
    set_para_spacing(p2, before=0, after=0)
    r2 = p2.add_run(school)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(8)
    r2.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)
    r2.italic = True
    p3 = left.add_paragraph()
    set_para_spacing(p3, before=0, after=60)
    r3 = p3.add_run(year)
    r3.font.name = 'Calibri'
    r3.font.size = Pt(8)
    r3.font.color.rgb = LBLUE

# — OTROS ESTUDIOS
sidebar_section(left, 'Otros Estudios')
for item in [
    'Diplomado en Transformación Digital\nUniv. de La Sabana (2019)',
    'Reportes GRI Standards G4\nGRI (2018, 2013, 2012)',
    'Starting Business\nShadd Business Centre, Canadá (2015)',
    'Inglés — Commission Scolaire\nEnglish-Montréal (2014-2015)',
    'Gestión ISO 9001:2008\nIcontec (2012)',
    'Auditor Ambiental en P+L\nIcontec (2010)',
]:
    p = left.add_paragraph()
    set_para_spacing(p, before=60, after=0)
    lines = item.split('\n')
    r1 = p.add_run(lines[0])
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(8)
    r1.font.color.rgb = WHITE
    if len(lines) > 1:
        p2 = left.add_paragraph()
        set_para_spacing(p2, before=0, after=40)
        r2 = p2.add_run(lines[1])
        r2.font.name = 'Calibri'
        r2.font.size = Pt(7.5)
        r2.font.color.rgb = RGBColor(0xAA, 0xBB, 0xCC)
        r2.italic = True

# ══════════════════════════════════════════════════════════════════════════
#  CONTENIDO PRINCIPAL  (columna derecha)
# ══════════════════════════════════════════════════════════════════════════

# — PERFIL
main_section(right, 'Perfil Profesional')
main_para(right, [
    ('Ingeniero, Magíster en Administración de Empresas y Transformación Exponencial, '
     'Especialista en Gerencia de Proyectos con ', False),
    ('más de 20 años de experiencia profesional ', True),
    ('en consultoría y dirección estratégica en los sectores minero, industrial y servicios, '
     'tanto en organizaciones gubernamentales como privadas.', False),
], before=60, after=60)
main_para(right, [
    ('Competencias en sostenibilidad corporativa, sistemas de gestión y modelos de '
     'transformación digital. Actualmente integra ', False),
    ('inteligencia artificial generativa (ChatGPT, Claude, Copilot) ', True),
    ('en procesos de reporte ESG, análisis de huella de carbono y automatización de '
     'flujos de trabajo de consultoría.', False),
], before=0, after=60)

# — EXPERIENCIA PROFESIONAL
main_section(right, 'Experiencia Profesional')

# 1. PPROJECTABILITY
job_header(right,
    'PPROJECTABILITY', 'Medellín, Colombia',
    'jun/2022 – Actualidad',
    'Director Estratégico')
sub_label(right, 'Funciones principales:')
for t in [
    'Estructuración de estrategias de sostenibilidad corporativa y adopción de enfoque ESG.',
    'Elaboración de estudios de materialidad, reportes de sostenibilidad y programas de gestión de emisiones.',
    'Coordinación de estrategias de relacionamiento con grupos de interés y adaptación al cambio climático.',
]:
    bullet_item(right, t)

sub_label(right, 'Logros destacados:')
for t in [
    'Diseño de planes de descarbonización aplicados en la industria petróleo y gas.',
    'Estructuración de soluciones digitales orientadas a estrategias de sostenibilidad.',
    'Materialización de proyectos de relacionamiento organizacional con grupos de interés.',
    'Integración de IA generativa (ChatGPT, Claude, Copilot) en flujos de consultoría ESG, '
     'optimizando la elaboración de reportes de sostenibilidad y estudios de materialidad.',
    'Análisis predictivo de huella de carbono con modelos de IA para escenarios de descarbonización '
     'en clientes del sector energético e industrial.',
    'Diseño de flujos y prompts con IA generativa para automatización de procesos internos '
     'y soluciones digitales para clientes corporativos.',
]:
    bullet_item(right, t)

# 2. PORTAFOLIO VERDE
job_header(right,
    'PORTAFOLIO VERDE', 'Medellín, Colombia',
    'jun/2015 – may/2022',
    'Ejecutivo de Proyectos / Consultor Senior')
sub_label(right, 'Funciones principales:')
for t in [
    'Diseño e implementación de planes de intervención para ejecución de proyectos de sostenibilidad.',
    'Coordinación de proyectos de gestión de huella de carbono organizacional.',
    'Principales clientes: Hocol, Stork, Celsia, Luker, Uniban, BBVA, Ministerio de Ambiente, Corantioquia.',
]:
    bullet_item(right, t)
sub_label(right, 'Logros:')
for t in [
    'Diseño y estructuración de metodologías de valoración empresarial en sostenibilidad.',
    'Estructuración de soluciones digitales dirigidas a estrategias de sostenibilidad corporativa.',
]:
    bullet_item(right, t)

# 3. ICONTEC
job_header(right,
    'ICONTEC', 'Medellín, Colombia',
    'ago/2012 – jul/2014',
    'Auditor en Sostenibilidad / Profesional de Evaluación de la Conformidad')
sub_label(right, 'Funciones y logros:')
for t in [
    'Evaluación de gestión en Responsabilidad Social e Informes de Sostenibilidad GRI.',
    'Diseño del "Sello de Sostenibilidad" y metodología de evaluación de mínimos para clientes.',
    'Clientes: Team Group, ESU, Mineros SA, Alcaldía de Medellín, Isagen.',
]:
    bullet_item(right, t)

# 4. CEMENTOS ARGOS
job_header(right,
    'CEMENTOS ARGOS', 'Medellín, Colombia',
    'ago/2010 – mar/2012',
    'Profesional de Gestión Ambiental')
for t in [
    'Gestión ambiental de instalaciones, monitoreo legal y operación de sistemas de gestión.',
    'Mejoramiento de indicadores legales ambientales mediante gestión diligente de cumplimiento normativo.',
]:
    bullet_item(right, t)

# 5. ÁREA METROPOLITANA
job_header(right,
    'ÁREA METROPOLITANA DEL VALLE ABURRÁ', 'Medellín, Colombia',
    'ene/2006 – jul/2010',
    'Consultor Ambiental')
for t in [
    'Formulación y seguimiento de proyectos de Producción Más Limpia para más de 500 empresas.',
    'Estructuración de herramientas de monitoreo y gestión de indicadores regionales de impacto.',
]:
    bullet_item(right, t)

# ══════════════════════════════════════════════════════════════════════════
#  GUARDAR
# ══════════════════════════════════════════════════════════════════════════
doc.save(OUTPUT)
print(f"CV generado: {OUTPUT}")
